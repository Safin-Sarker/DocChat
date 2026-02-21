"""DOCX extraction utilities."""

import logging
from typing import Any, Dict, List

from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


def _rows_to_markdown(rows: List[List[str]]) -> str:
    """Convert table rows to markdown text."""
    if not rows:
        return ""

    column_count = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (column_count - len(row)) for row in rows]
    header = normalized_rows[0]
    body = normalized_rows[1:]

    header_line = "| " + " | ".join(cell or "" for cell in header) + " |"
    sep_line = "| " + " | ".join("---" for _ in header) + " |"
    body_lines = [
        "| " + " | ".join(cell or "" for cell in row) + " |"
        for row in body
    ]
    return "\n".join([header_line, sep_line] + body_lines)


class DocxExtractor:
    """Extract text and tables from DOCX files."""

    def extract_pages(self, docx_path: str) -> List[Dict[str, Any]]:
        """Extract paragraph text into page-like sections."""
        try:
            document = DocxDocument(docx_path)
        except Exception as exc:
            logger.error("DOCX load failed: %s", exc)
            return []

        pages: List[Dict[str, Any]] = []
        current_lines: List[str] = []
        section_num = 1

        for paragraph in document.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue

            style_name = (paragraph.style.name if paragraph.style else "").lower()
            is_heading = style_name.startswith("heading")
            if is_heading and current_lines:
                pages.append({"page_num": section_num, "text": "\n".join(current_lines)})
                section_num += 1
                current_lines = [text]
            else:
                current_lines.append(text)

        if current_lines:
            pages.append({"page_num": section_num, "text": "\n".join(current_lines)})

        return pages

    def extract_tables(self, docx_path: str) -> List[Dict[str, Any]]:
        """Extract DOCX tables as markdown entries."""
        try:
            document = DocxDocument(docx_path)
        except Exception as exc:
            logger.error("DOCX load failed for tables: %s", exc)
            return []

        tables_out: List[Dict[str, Any]] = []
        for table_index, table in enumerate(document.tables):
            rows: List[List[str]] = []
            for row in table.rows:
                rows.append([(cell.text or "").strip() for cell in row.cells])
            markdown = _rows_to_markdown(rows)
            if not markdown:
                continue

            tables_out.append({
                "page": 1,
                "table_index": table_index,
                "markdown": markdown,
            })

        return tables_out
